import streamlit as st
import fitz
import pytesseract
from pdf2image import convert_from_bytes
from docx import Document as DocxReader
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from groq import Groq
import io
import datetime


# ─── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="LMS Content Generator",
    page_icon="🎓",
    layout="centered"
)


# ─── Header ───────────────────────────────────────────────────────────────────
st.title("🎓 LMS Content Generator")
st.caption("Upload product documents → Get a structured training module draft")
st.divider()


# ─── API Key (from Streamlit Secrets — set once, hidden from users) ────────────
GROQ_API_KEY = st.secrets["GROQ_API_KEY"]


# ─── System Prompt ────────────────────────────────────────────────────────────
SYSTEM_PROMPT = """
You are an expert learning content designer creating training modules for field
sales staff and ground-level sales representatives at an FMCG company.


Your task is to take raw product source documents and transform them into a
structured, easy-to-understand training module.


STRICT OUTPUT RULES:
- Always generate exactly the 6 sections listed below, in order.
- Do not add any extra sections, preamble, or closing remarks.
- Use simple, direct language. Reading level: Class 8 to 10 equivalent.
- Avoid technical jargon unless absolutely necessary.


═══════════════════════════════════════════════
SECTION 1: PRODUCT OVERVIEW
═══════════════════════════════════════════════
- Write 3 to 4 sentences.
- Cover: what the product is, who it is for, and what makes it relevant.


═══════════════════════════════════════════════
SECTION 2: KEY FEATURES (max 5 bullet points)
═══════════════════════════════════════════════
- Each bullet is one sentence maximum.
- Focus on what a salesperson needs to communicate.
- Format: start each bullet with a dash ( - )


═══════════════════════════════════════════════
SECTION 3: CUSTOMER PAIN POINTS ADDRESSED
═══════════════════════════════════════════════
- List 3 bullet points.
- Format: PROBLEM: [problem] / SOLUTION: [solution]


═══════════════════════════════════════════════
SECTION 4: SALES PITCH FRAMEWORK
═══════════════════════════════════════════════
OPENING LINE: [one sentence to introduce the product]
VALUE POINT 1: [key benefit]
VALUE POINT 2: [second benefit]
VALUE POINT 3: [third benefit]
CLOSING LINE: [one sentence to drive purchase]


═══════════════════════════════════════════════
SECTION 5: COMMON OBJECTIONS AND RESPONSES
═══════════════════════════════════════════════
- List 3 objections.
- Format: OBJECTION: [objection] / RESPONSE: [response]


═══════════════════════════════════════════════
SECTION 6: KNOWLEDGE CHECK
═══════════════════════════════════════════════
- 5 multiple choice questions.
- Format each as:
  Q[number]: [question]
  A) B) C) D) options
  CORRECT ANSWER: [letter]
"""


# ─── Helper functions ─────────────────────────────────────────────────────────
def extract_text_from_pdf(file_bytes):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page_num, page in enumerate(doc):
        page_text = page.get_text()
        if page_text.strip():
            text += f"[Page {page_num + 1}]\n{page_text}\n"
    return text


def extract_text_from_pdf_ocr(file_bytes):
    images = convert_from_bytes(file_bytes, dpi=200)
    ocr_text = ""
    for i, image in enumerate(images):
        page_text = pytesseract.image_to_string(image, lang="eng")
        if page_text.strip():
            ocr_text += f"[Page {i+1}]\n{page_text}\n\n"
    return ocr_text


def extract_text_from_docx(file_bytes):
    doc = DocxReader(io.BytesIO(file_bytes))
    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text += row_text + "\n"
    return text


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_font(run, size_pt, bold=False, italic=False, color_hex=None):
    run.font.name = 'Arial'
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color_hex:
        r, g, b = int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16)
        run.font.color.rgb = RGBColor(r, g, b)


def build_word_doc(generated_module, source_filenames):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)


    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Product Training Module")
    set_font(title_run, 22, bold=True, color_hex="1A1915")


    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(f"Sources: {', '.join(source_filenames)}")
    set_font(meta_run, 9, italic=True, color_hex="7A7260")
    date_run = meta_para.add_run(f"     Generated: {datetime.datetime.now().strftime('%d %B %Y')}")
    set_font(date_run, 9, italic=True, color_hex="7A7260")


    draft_para = doc.add_paragraph()
    draft_run = draft_para.add_run("⚠  DRAFT — AI-generated. SDM review required before LMS upload.")
    set_font(draft_run, 9, bold=True, color_hex="C0392B")


    add_horizontal_rule(doc)
    doc.add_paragraph()


    SECTION_KEYWORDS = [
        "SECTION 1","SECTION 2","SECTION 3","SECTION 4","SECTION 5","SECTION 6",
        "PRODUCT OVERVIEW","KEY FEATURES","CUSTOMER PAIN POINTS",
        "SALES PITCH","COMMON OBJECTIONS","KNOWLEDGE CHECK"
    ]
    FIELD_LABELS = [
        "PROBLEM:","SOLUTION:","OBJECTION:","RESPONSE:",
        "OPENING LINE:","VALUE POINT","CLOSING LINE:","CORRECT ANSWER:"
    ]


    for line in generated_module.strip().split("\n"):
        stripped = line.strip()
        if stripped == "":
            doc.add_paragraph("")
            continue
        is_section_header = any(kw in stripped.upper() for kw in SECTION_KEYWORDS)
        if is_section_header and (stripped.isupper() or stripped.upper().startswith("SECTION")):
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(stripped.replace("═","").strip())
            set_font(run, 12, bold=True, color_hex="1A1915")
            add_horizontal_rule(doc)
            continue
        if all(c in "═─=─ " for c in stripped) and len(stripped) > 3:
            continue
        if stripped.upper().startswith("Q") and len(stripped) > 2 and stripped[1].isdigit() and ":" in stripped[:5]:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            set_font(run, 10, bold=True, color_hex="1A1915")
            continue
        if len(stripped) > 2 and stripped[0].upper() in "ABCD" and stripped[1] == ")" and stripped[2] == " ":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(stripped)
            set_font(run, 10, color_hex="2C2C2A")
            continue
        if stripped.upper().startswith("CORRECT ANSWER:"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(stripped)
            set_font(run, 10, bold=True, color_hex="1F7A4A")
            continue
        is_field_label = any(stripped.upper().startswith(label) for label in FIELD_LABELS)
        if is_field_label:
            colon_pos = stripped.find(":")
            if colon_pos != -1:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                label_run = p.add_run(stripped[:colon_pos+1])
                set_font(label_run, 10, bold=True, color_hex="3D3A32")
                value_run = p.add_run(stripped[colon_pos+1:])
                set_font(value_run, 10, color_hex="1A1915")
                continue
        if stripped.startswith("- ") or stripped.startswith("• "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(stripped[2:].strip())
            set_font(run, 10, color_hex="1A1915")
            continue
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        set_font(run, 10, color_hex="1A1915")


    doc.add_paragraph()
    add_horizontal_rule(doc)
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        "Generated by LMS Content Pipeline using Groq (Llama 3.3). "
        "Review for accuracy before uploading to LMS."
    )
    set_font(footer_run, 8, italic=True, color_hex="A09880")


    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ─── Main UI ──────────────────────────────────────────────────────────────────
uploaded_files = st.file_uploader(
    "Upload product documents",
    type=["pdf", "docx"],
    accept_multiple_files=True,
    help="Upload 1 to 3 files — brand manager brief, spec sheet, product note, etc."
)


if uploaded_files:
    st.success(f"{len(uploaded_files)} file(s) uploaded")


    if st.button("Generate Training Module", type="primary", use_container_width=True):


        combined_text = ""


        with st.status("Processing documents...", expanded=True) as status:


            for uploaded_file in uploaded_files:
                st.write(f"Extracting: {uploaded_file.name}...")
                file_bytes = uploaded_file.read()


                if uploaded_file.name.lower().endswith(".pdf"):
                    text = extract_text_from_pdf(file_bytes)
                    if len(text.strip()) < 100:
                        st.write(f"→ No text layer, running OCR on {uploaded_file.name}...")
                        text = extract_text_from_pdf_ocr(file_bytes)
                elif uploaded_file.name.lower().endswith(".docx"):
                    text = extract_text_from_docx(file_bytes)


                combined_text += f"\n\n{'='*60}\nSOURCE: {uploaded_file.name}\n{'='*60}\n\n{text}"
                st.write(f"✅ {uploaded_file.name} — {len(text):,} characters extracted")


            st.write("Sending to Groq AI...")
            client = Groq(api_key=GROQ_API_KEY)
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": f"Here are the product documents. Generate the training module.\n\n{combined_text}"}
                ],
                temperature=0.4,
                max_tokens=3500,
            )
            generated_module = response.choices[0].message.content
            st.write("✅ Module generated")


            st.write("Building Word document...")
            source_names = [f.name for f in uploaded_files]
            word_buffer = build_word_doc(generated_module, source_names)
            status.update(label="Done!", state="complete")


        st.divider()
        st.subheader("Module Preview")
        st.text_area("Generated content", generated_module, height=400)


        st.download_button(
            label="⬇️  Download training_module_draft.docx",
            data=word_buffer,
            file_name="training_module_draft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary"
        )